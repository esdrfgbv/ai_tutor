"""
OCR Service
=============
Text extraction from documents with OCR fallback for scanned content.
Extracts text from images embedded within PDF pages so no content is lost.
Supports native PDFs, scanned PDFs, DOCX, TXT, and images.
"""

from __future__ import annotations

import logging
import os
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF

from app.models.enums import DocumentType

logger = logging.getLogger(__name__)

# ── Tesseract initialisation ───────────────────────────────────────────────
_TESSERACT_PATHS = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
]


@dataclass
class PageContent:
    """Extracted content from a single page."""
    page_number: int
    text: str
    is_ocr: bool = False
    width: float = 0.0
    height: float = 0.0
    images: list[dict] = field(default_factory=list)


@dataclass
class ExtractionResult:
    """Result of text extraction from a document."""
    pages: list[PageContent]
    total_pages: int
    ocr_pages: int = 0
    native_text_pages: int = 0
    extraction_method: str = "pymupdf"
    errors: list[str] = field(default_factory=list)


class OCRService:
    """Handles text extraction from all supported document types."""

    def __init__(self):
        self._has_ocr = self._check_ocr_available()
        if self._has_ocr:
            self._setup_tesseract()

    @staticmethod
    def _setup_tesseract():
        """Point pytesseract to the installed Tesseract binary."""
        import pytesseract
        for tp in _TESSERACT_PATHS:
            if os.path.exists(tp):
                pytesseract.pytesseract.tesseract_cmd = tp
                break

    def extract(self, file_path: Path, document_type: DocumentType) -> ExtractionResult:
        """
        Main entry point. Routes to the appropriate extractor based on document type.
        """
        if document_type == DocumentType.pdf:
            return self._extract_pdf(file_path)
        elif document_type == DocumentType.docx:
            return self._extract_docx(file_path)
        elif document_type == DocumentType.image:
            return self._extract_image(file_path)
        elif document_type == DocumentType.txt:
            return self._extract_txt(file_path)
        else:
            return ExtractionResult(
                pages=[], total_pages=0,
                errors=[f"Unsupported document type: {document_type}"],
            )

    # ── PDF Extraction ───────────────────────────────────────────────────

    def _extract_pdf(self, file_path: Path) -> ExtractionResult:
        """Extract text from PDF using PyMuPDF with OCR fallback for every image."""
        pages: list[PageContent] = []
        ocr_count = 0
        native_count = 0
        errors: list[str] = []

        try:
            doc = fitz.open(str(file_path))
        except Exception as e:
            return ExtractionResult(
                pages=[], total_pages=0,
                errors=[f"Cannot open PDF: {e}"],
            )

        for idx, page in enumerate(doc, start=1):
            page_width = page.rect.width
            page_height = page.rect.height

            # 1. Native text
            text = page.get_text("text")

            # 2. OCR every embedded image and append its text
            image_list: list[dict] = []
            image_ocr_texts: list[str] = []
            try:
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    try:
                        base_image = doc.extract_image(xref)
                        if not base_image:
                            continue
                        w = base_image.get("width", 0)
                        h = base_image.get("height", 0)
                        if w < 30 or h < 30:
                            continue
                        image_list.append({
                            "xref": xref,
                            "width": w,
                            "height": h,
                            "ext": base_image.get("ext", "png"),
                        })
                        # OCR the image bytes
                        if self._has_ocr:
                            img_bytes = base_image.get("image")
                            if img_bytes:
                                ocr_img_text = self._ocr_image_bytes(img_bytes)
                                if ocr_img_text:
                                    image_ocr_texts.append(ocr_img_text)
                    except Exception:
                        pass
            except Exception:
                pass

            is_ocr = False

            # 3. If page has very little native text, OCR the rendered page
            native_text_len = len((text or "").strip())
            if native_text_len < 50 and self._has_ocr:
                ocr_text = self._ocr_page(page)
                if ocr_text and len(ocr_text.strip()) > native_text_len:
                    text = ocr_text
                    is_ocr = True
                    ocr_count += 1
                elif text and text.strip():
                    native_count += 1
            elif text and text.strip():
                native_count += 1

            # 4. Append any text found inside embedded images
            if image_ocr_texts:
                separator = "\n---[Image text]---\n" if text.strip() else ""
                text = (text or "") + separator + "\n".join(image_ocr_texts)

            if text and text.strip():
                pages.append(PageContent(
                    page_number=idx,
                    text=text,
                    is_ocr=is_ocr,
                    width=page_width,
                    height=page_height,
                    images=image_list,
                ))

        doc.close()

        return ExtractionResult(
            pages=pages,
            total_pages=len(pages),
            ocr_pages=ocr_count,
            native_text_pages=native_count,
            extraction_method="pymupdf" + ("+ocr" if ocr_count > 0 else ""),
            errors=errors,
        )

    def _ocr_image_bytes(self, image_bytes: bytes) -> str:
        """Run OCR on raw image bytes (JPEG/PNG from inside a PDF)."""
        try:
            import pytesseract
            from PIL import Image
            import io

            img = Image.open(io.BytesIO(image_bytes))
            img = self._preprocess_image(img)
            # Assume single uniform block of text
            text = pytesseract.image_to_string(
                img, lang="eng+hin", config="--psm 6",
            )
            return text.strip()
        except Exception as e:
            logger.debug("Image OCR failed: %s", e)
            return ""

    @staticmethod
    def _preprocess_image(img):
        """Preprocess a PIL image for better OCR accuracy."""
        from PIL import ImageFilter, ImageEnhance

        if img.mode != "L":
            img = img.convert("L")

        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)
        img = img.filter(ImageFilter.SHARPEN)

        return img

    def _ocr_page(self, page) -> str:
        """Run OCR on a single PDF page at high resolution."""
        try:
            import pytesseract
            from PIL import Image
            import io

            # Render page at 2x resolution for OCR
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            img = self._preprocess_image(img)
            # Automatic page segmentation (handles multi-column layouts)
            text = pytesseract.image_to_string(
                img, lang="eng+hin", config="--psm 3",
            )
            return text
        except Exception as e:
            logger.warning("Page OCR failed: %s", e)
            return ""

    def _check_ocr_available(self) -> bool:
        """Check if OCR dependencies are available."""
        try:
            import pytesseract  # noqa: F401
            from PIL import Image  # noqa: F401
            return True
        except ImportError:
            logger.info("pytesseract/PIL not installed. OCR disabled.")
            return False

    # ── DOCX Extraction ──────────────────────────────────────────────────

    def _extract_docx(self, file_path: Path) -> ExtractionResult:
        """Extract text from DOCX files."""
        try:
            import docx
        except ImportError:
            return ExtractionResult(
                pages=[], total_pages=0,
                errors=["python-docx not installed. Cannot process DOCX files."],
            )

        try:
            document = docx.Document(str(file_path))
        except Exception as e:
            return ExtractionResult(
                pages=[], total_pages=0,
                errors=[f"Cannot open DOCX: {e}"],
            )

        # DOCX doesn't have pages, so we treat the whole document as page 1
        # but split by sections/headings for better chunking later
        full_text_parts: list[str] = []
        for para in document.paragraphs:
            if para.text.strip():
                full_text_parts.append(para.text)

        # Also extract tables as text
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    full_text_parts.append(" | ".join(cells))

        full_text = "\n".join(full_text_parts)

        pages = []
        if full_text.strip():
            pages.append(PageContent(
                page_number=1,
                text=full_text,
                is_ocr=False,
            ))

        return ExtractionResult(
            pages=pages,
            total_pages=1 if pages else 0,
            extraction_method="python-docx",
        )

    # ── Image Extraction ─────────────────────────────────────────────────

    def _extract_image(self, file_path: Path) -> ExtractionResult:
        """Extract text from images using OCR."""
        if not self._has_ocr:
            return ExtractionResult(
                pages=[], total_pages=0,
                errors=["OCR not available for image processing."],
            )

        try:
            import pytesseract
            from PIL import Image

            img = Image.open(str(file_path))
            img = self._preprocess_image(img)
            text = pytesseract.image_to_string(img, lang="eng+hin")

            pages = []
            if text and text.strip():
                pages.append(PageContent(
                    page_number=1,
                    text=text,
                    is_ocr=True,
                    width=float(img.width),
                    height=float(img.height),
                ))

            return ExtractionResult(
                pages=pages,
                total_pages=1 if pages else 0,
                ocr_pages=1 if pages else 0,
                extraction_method="pytesseract",
            )
        except Exception as e:
            return ExtractionResult(
                pages=[], total_pages=0,
                errors=[f"Image OCR failed: {e}"],
            )

    # ── TXT Extraction ───────────────────────────────────────────────────

    def _extract_txt(self, file_path: Path) -> ExtractionResult:
        """Read plain text files."""
        try:
            text = file_path.read_text(encoding="utf-8", errors="replace")
            pages = []
            if text.strip():
                pages.append(PageContent(
                    page_number=1,
                    text=text,
                    is_ocr=False,
                ))
            return ExtractionResult(
                pages=pages,
                total_pages=1 if pages else 0,
                extraction_method="plaintext",
            )
        except Exception as e:
            return ExtractionResult(
                pages=[], total_pages=0,
                errors=[f"Cannot read text file: {e}"],
            )


ocr_service = OCRService()
